# -*- coding: utf-8 -*-
import requests
import os
import json
import time
from requests.exceptions import RequestException
import subprocess
import sys
import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog
import threading
import webbrowser
import queue
import sqlite3
from datetime import datetime

# 配置信息：定义了调用Coze API的地址、认证Token、两个不同机器人的ID、用户ID以及超时时间
AGENT_API_URL = "https://api.coze.cn/v3/chat"
AGENT_TOKEN = "pat_DUpJqpypgHeEV904AQpVoLOYPK8oCOBr0oAlPuz96V5Dko6E4Nt2baBKZCuxKqRD"
OLD_BOT_ID = "7547920557141409807"
NEW_BOT_ID = "7572591876910596148"
USER_ID = "english_test_user_001"
OLD_TIMEOUT = 200
NEW_TIMEOUT = 200

# 数据库管理类：负责创建、保存、查询和删除试卷记录，使用SQLite存储
class ExamDatabase:
    # 初始化数据库文件并创建表格（如果不存在）
    def __init__(self, db_path="exam_database.db"):
        self.db_path = db_path
        self.init_database()

    # 创建试卷信息表，包含标题、各个部分原文、题目答案、文件路径和创建时间等字段
    def init_database(self):
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS exams (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT NOT NULL,
                essay_topic TEXT,
                cloze_text TEXT,
                reading_a TEXT,
                reading_b TEXT,
                reading_c TEXT,
                question_content TEXT,
                answer_content TEXT,
                html_path TEXT,
                word_path TEXT,
                md_path TEXT,
                created_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        conn.commit()
        conn.close()

    # 将生成的试卷信息保存到数据库，返回新插入的记录ID
    def save_exam(self, title, essay_topic, cloze_text, reading_a, reading_b, reading_c,
                  question_content, answer_content, html_path, word_path, md_path):
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO exams
            (title, essay_topic, cloze_text, reading_a, reading_b, reading_c,
             question_content, answer_content, html_path, word_path, md_path)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (title, essay_topic, cloze_text, reading_a, reading_b, reading_c,
              question_content, answer_content, html_path, word_path, md_path))
        exam_id = cursor.lastrowid
        conn.commit()
        conn.close()
        return exam_id

    # 获取所有试卷记录，按创建时间倒序排列，仅返回ID、标题和时间
    def get_all_exams(self):
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('SELECT id, title, created_time FROM exams ORDER BY created_time DESC')
        exams = cursor.fetchall()
        conn.close()
        return exams

    # 根据ID查询单条完整试卷记录
    def get_exam_by_id(self, exam_id):
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM exams WHERE id = ?', (exam_id,))
        exam = cursor.fetchone()
        conn.close()
        return exam

    # 根据ID删除试卷记录
    def delete_exam(self, exam_id):
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('DELETE FROM exams WHERE id = ?', (exam_id,))
        conn.commit()
        conn.close()

# 主GUI类：负责整个应用程序的界面布局、交互逻辑和试卷生成流程控制
class EnglishExamGeneratorGUI:
    # 初始化主窗口，设置标题、尺寸、背景，并创建日志队列和数据库实例
    def __init__(self, root):
        self.root = root
        self.root.title("英语试卷生成器")
        self.root.geometry("1300x900")
        self.root.configure(bg='#f0f0f0')
        self.log_queue = queue.Queue()
        self.is_generating = False
        self.generated_files = []
        self.db = ExamDatabase()
        self.setup_ui()
        self.update_log()

    # 构建完整的界面布局，包括输入区、按钮、日志区和进度条
    def setup_ui(self):
        tk.Label(self.root, text="英语试卷生成器", font=("黑体", 26, "bold"), bg='#f0f0f0', fg='#2c3e50').pack(pady=20)
        main_container = ttk.Frame(self.root)
        main_container.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 10))
        left_container = ttk.Frame(main_container)
        left_container.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        right_frame = ttk.Frame(main_container)
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(10, 0))

        # 左侧可滚动输入区域，用于输入作文、完形、阅读等原文
        canvas = tk.Canvas(left_container, highlightthickness=0)
        scrollbar = ttk.Scrollbar(left_container, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        input_frame = ttk.LabelFrame(scrollable_frame, text="试卷内容设置", padding=15)
        input_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        # 创建统一的输入段落组件（标签+文本框+导入按钮）
        def create_section(label):
            frame = ttk.Frame(input_frame)
            frame.pack(fill=tk.X, pady=12)
            tk.Label(frame, text=label, font=("黑体", 11, "bold"), anchor="w").pack(fill=tk.X)
            text = scrolledtext.ScrolledText(frame, height=5, font=("宋体", 10), wrap=tk.WORD)
            text.pack(fill=tk.X, pady=(2, 5))
            ttk.Button(frame, text="导入文件", command=lambda: self.import_file_to_text(text)).pack(side=tk.RIGHT)
            return text

        self.essay_text = create_section("请输入作文题材：")
        self.cloze_text = create_section("请输入完形填空原文：")
        self.reading_a_text = create_section("请输入阅读A原文：")
        self.reading_b_text = create_section("请输入阅读B原文：")
        self.reading_c_text = create_section("请输入阅读C原文：")

        # 底部按钮区域
        bottom_frame = ttk.Frame(self.root)
        bottom_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=20)
        self.generate_btn = ttk.Button(bottom_frame, text="开始生成试卷", command=self.start_generation, style="Accent.TButton")
        self.generate_btn.pack(side=tk.LEFT, padx=40)
        self.open_files_btn = ttk.Button(bottom_frame, text="一键打开Word文档", command=self.open_word_document, state="disabled")
        self.open_files_btn.pack(side=tk.LEFT, padx=20)
        self.history_btn = ttk.Button(bottom_frame, text="查看历史试卷", command=self.show_history)
        self.history_btn.pack(side=tk.LEFT, padx=20)

        # 右侧日志和进度显示区域
        log_frame = ttk.LabelFrame(right_frame, text="生成过程日志", padding=10)
        log_frame.pack(fill=tk.BOTH, expand=True)
        self.log_text = scrolledtext.ScrolledText(log_frame, font=("Consolas", 9), wrap=tk.WORD, state="disabled")
        self.log_text.pack(fill=tk.BOTH, expand=True)
        self.progress = ttk.Progressbar(right_frame, mode='indeterminate')
        self.progress.pack(fill=tk.X, pady=10)
        self.status_label = tk.Label(right_frame, text="就绪", font=("宋体", 10), bg='#f0f0f0', fg='#666666')
        self.status_label.pack()

        # 绑定鼠标滚轮事件实现滚动
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", on_mousewheel)

    # 弹出文件选择对话框，将选中的txt或docx文件内容导入到指定的文本框
    def import_file_to_text(self, text_widget):
        path = filedialog.askopenfilename(
            title="选择文件导入",
            filetypes=[("支持格式", "*.txt *.docx *.md *.html"), ("所有文件", "*.*")]
        )
        if not path: return
        content = self.read_file(path)
        if content:
            text_widget.delete(1.0, tk.END)
            text_widget.insert(tk.END, content.strip())

    # 读取文件内容，支持txt和docx格式，docx会自动尝试安装python-docx依赖
    def read_file(self, path):
        try:
            if path.lower().endswith(".docx"):
                try:
                    from docx import Document
                    doc = Document(path)
                    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                except ImportError:
                    messagebox.showinfo("安装依赖", "正在安装 python-docx...")
                    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
                    from docx import Document
                    doc = Document(path)
                    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            else:
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()
        except Exception as e:
            messagebox.showerror("错误", f"读取失败：{e}")
            return None

    # 将消息放入日志队列，供界面线程定期刷新显示
    def log_message(self, message):
        self.log_queue.put(message)

    # 定时从队列中取出日志消息并追加到日志文本框，实现实时更新
    def update_log(self):
        try:
            while True:
                msg = self.log_queue.get_nowait()
                self.log_text.configure(state="normal")
                self.log_text.insert(tk.END, msg + "\n")
                self.log_text.see(tk.END)
                self.log_text.configure(state="disabled")
        except queue.Empty:
            pass
        finally:
            self.root.after(100, self.update_log)

    # 更新底部状态栏文字
    def update_status(self, message):
        self.status_label.config(text=message)
        self.root.update_idletasks()

    # 开始生成试卷的入口函数，禁用按钮并启动后台线程
    def start_generation(self):
        if self.is_generating: return
        self.is_generating = True
        self.generate_btn.config(state="disabled")
        self.open_files_btn.config(state="disabled")
        self.history_btn.config(state="disabled")
        self.progress.start(10)
        self.log_text.configure(state="normal")
        self.log_text.delete(1.0, tk.END)
        self.log_text.configure(state="disabled")
        threading.Thread(target=self.generate_exam, daemon=True).start()

    # 核心生成流程：在后台线程中依次调用两个AI机器人生成试卷和答案，并保存文件及数据库记录
    def generate_exam(self):
        try:
            self.log_message("=" * 50)
            self.log_message("开始生成英语试卷...")
            self.log_message("=" * 50)
            exam_prompt = self.build_exam_prompt()
            self.log_message(f"\n发送给智能体的提示：\n{exam_prompt}")
            self.update_status("检查依赖包...")
            self.log_message("\n1. 检查依赖包...")
            has_dependencies = check_and_install_dependencies()
            self.update_status("调用试卷生成器...")
            self.log_message("\n2. 调用试卷生成器生成完整英语试卷...")
            old_content = call_agent_stream_gui(exam_prompt, OLD_BOT_ID, OLD_TIMEOUT, self, "试卷生成器")
            if not old_content:
                self.log_message("试卷生成器未生成有效内容，终止流程")
                return
            self.update_status("调用答案分离器...")
            self.log_message("\n3. 调用答案分离器处理内容...")
            new_prompt = f"处理以下英语试卷内容：\n{old_content}"
            new_content = call_agent_stream_gui(new_prompt, NEW_BOT_ID, NEW_TIMEOUT, self, "答案分离器")
            if not new_content:
                self.log_message("答案分离器未返回有效内容")
                return

            # 去除重复行
            lines = new_content.splitlines()
            unique_lines = []
            seen = set()
            for line in lines:
                if line not in seen:
                    seen.add(line)
                    unique_lines.append(line)
            new_content = "\n".join(unique_lines)

            # 从返回内容中提取题目和答案部分的代码块
            question_content = extract_code_block_content(new_content, "题目")
            answer_content = extract_code_block_content(new_content, "答案")
            question_content = question_content or ""
            answer_content = answer_content or ""

            self.generated_files = []
            html_path = self.generate_html_files(question_content, answer_content)
            if html_path: self.generated_files.append(html_path)

            word_path = "english_comprehensive_exam.docx"
            if has_dependencies and markdown_to_word(question_content, answer_content, word_path):
                self.generated_files.append(word_path)

            md_path = "english_comprehensive_exam.md"
            with open(md_path, "w", encoding="utf-8") as f:
                f.write("# 英语综合试卷题目\n\n" + question_content + "\n\n# 参考答案与解析\n\n" + answer_content)
            self.generated_files.append(md_path)

            # 生成标题并保存到数据库
            title = f"英语试卷_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            self.db.save_exam(
                title=title,
                essay_topic=self.essay_text.get(1.0, tk.END).strip(),
                cloze_text=self.cloze_text.get(1.0, tk.END).strip(),
                reading_a=self.reading_a_text.get(1.0, tk.END).strip(),
                reading_b=self.reading_b_text.get(1.0, tk.END).strip(),
                reading_c=self.reading_c_text.get(1.0, tk.END).strip(),
                question_content=question_content,
                answer_content=answer_content,
                html_path=html_path if html_path else "",
                word_path=word_path if word_path and os.path.exists(word_path) else "",
                md_path=md_path
            )

            self.log_message("试卷生成完成！")
            self.open_files_btn.config(state="normal")
            messagebox.showinfo("完成", "英语试卷生成完成！")
        except Exception as e:
            self.log_message(f"\n错误：{str(e)}")
            messagebox.showerror("错误", f"生成过程中出现错误：{str(e)}")
        finally:
            self.is_generating = False
            self.generate_btn.config(state="normal")
            self.history_btn.config(state="normal")
            self.progress.stop()

    # 根据用户输入或缺省值构建发送给第一个AI的完整提示
    def build_exam_prompt(self):
        essay = self.essay_text.get(1.0, tk.END).strip()
        cloze = self.cloze_text.get(1.0, tk.END).strip()
        reading_a = self.reading_a_text.get(1.0, tk.END).strip()
        reading_b = self.reading_b_text.get(1.0, tk.END).strip()
        reading_c = self.reading_c_text.get(1.0, tk.END).strip()
        parts = []
        parts.append(f"作文题材：{essay}" if essay else "作文题材：请自行随机生成")
        parts.append(f"完形填空原文：{cloze}" if cloze else "完形填空原文：请自行随机生成")
        parts.append(f"阅读A原文：{reading_a}" if reading_a else "阅读A原文：请自行随机生成")
        parts.append(f"阅读B原文：{reading_b}" if reading_b else "阅读B原文：请自行随机生成")
        parts.append(f"阅读C原文：{reading_c}" if reading_c else "阅读C原文：请自行随机生成")
        return "\n".join(parts)

    # 根据题目和答案内容生成带样式的HTML文件（包含题目和答案两部分）
    def generate_html_files(self, question_content, answer_content):
        try:
            html_header = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>英语综合试卷</title>
    <style>
        body { font-family: Arial, 'SimHei', 'Microsoft YaHei', sans-serif; line-height: 1.4; padding: 30px; max-width: 900px; margin: 0 auto; background-color: #f8f9fa; }
        .section { background: white; padding: 20px; margin: 15px 0; border-radius: 8px; box-shadow: 0 2px 5px rgba(0,0,0,0.1); }
        .answer-section { margin-top: 30px; padding-top: 20px; border-top: 2px solid #3498db; background: #f8f9fa; }
        table { width: 100%; border-collapse: collapse; margin: 10px 0; }
        th, td { border: 1px solid #ddd; padding: 10px; text-align: left; }
        th { background-color: #2c3e50; color: white; font-weight: bold; }
        tr:nth-child(even) { background-color: #f2f2f2; }
        .option { margin: 3px 0 3px 25px; }
    </style>
</head>
<body>
    <h1 style='color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 15px; text-align: center; margin-bottom: 20px;'>英语综合能力测试试卷</h1>
"""
            formatted_question = format_content_to_html(question_content) if question_content else ""
            question_html = f"<div class='section'><h2 style='color: #2980b9; text-align: center;'>试卷题目</h2>\n{formatted_question}</div>" if question_content else ""
            formatted_answer = format_content_to_html(answer_content) if answer_content else ""
            answer_html = f"<div class='section answer-section'><h2 style='color: #e74c3c; text-align: center;'>参考答案与解析</h2>\n{formatted_answer}</div>" if answer_content else ""
            html_footer = "</body></html>"
            path = "english_comprehensive_exam.html"
            with open(path, "w", encoding="utf-8") as f:
                f.write(html_header + question_html + answer_html + html_footer)
            return path
        except Exception as e:
            self.log_message(f"生成HTML文件失败: {str(e)}")
            return None

    # 打开当前生成的Word文档（使用系统默认程序）
    def open_word_document(self):
        word_path = "english_comprehensive_exam.docx"
        if os.path.exists(word_path):
            os.startfile(os.path.abspath(word_path))
        else:
            messagebox.showwarning("警告", f"未找到Word文档: {word_path}")

    # 打开历史试卷查看窗口，显示试卷列表和详情
    def show_history(self):
        history_window = tk.Toplevel(self.root)
        history_window.title("历史试卷")
        history_window.geometry("1000x700")
        list_frame = ttk.Frame(history_window)
        list_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        left_frame = ttk.Frame(list_frame)
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        right_frame = ttk.Frame(list_frame)
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(10, 0))

        ttk.Label(left_frame, text="历史试卷列表", font=("黑体", 12, "bold")).pack(anchor="w")
        self.history_listbox = tk.Listbox(left_frame)
        self.history_listbox.pack(fill=tk.BOTH, expand=True, pady=5)
        self.history_listbox.bind('<<ListboxSelect>>', self.on_history_select)
        ttk.Button(left_frame, text="刷新列表", command=self.refresh_history).pack(fill=tk.X)

        ttk.Label(right_frame, text="试卷详情", font=("黑体", 12, "bold")).pack(anchor="w")
        notebook = ttk.Notebook(right_frame)
        notebook.pack(fill=tk.BOTH, expand=True, pady=5)

        question_frame = ttk.Frame(notebook)
        self.question_text = scrolledtext.ScrolledText(question_frame, wrap=tk.WORD)
        self.question_text.pack(fill=tk.BOTH, expand=True)
        notebook.add(question_frame, text="题目")

        answer_frame = ttk.Frame(notebook)
        self.answer_text = scrolledtext.ScrolledText(answer_frame, wrap=tk.WORD)
        self.answer_text.pack(fill=tk.BOTH, expand=True)
        notebook.add(answer_frame, text="答案")

        button_frame = ttk.Frame(right_frame)
        button_frame.pack(fill=tk.X, pady=5)
        ttk.Button(button_frame, text="打开Word文档", command=self.open_selected_word).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(button_frame, text="打开HTML文件", command=self.open_selected_html).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(button_frame, text="删除选中试卷", command=self.delete_selected_exam).pack(side=tk.LEFT)

        self.refresh_history()

    # 刷新历史试卷列表，从数据库读取并更新Listbox显示
    def refresh_history(self):
        self.history_listbox.delete(0, tk.END)
        exams = self.db.get_all_exams()
        self.exam_mapping = {}
        for i, (exam_id, title, created_time) in enumerate(exams):
            time_str = created_time.split('.')[0]
            display_text = f"{exam_id}: {title} ({time_str})"
            self.history_listbox.insert(tk.END, display_text)
            self.exam_mapping[i] = exam_id

    # 当用户在列表中选择某条试卷时，加载对应的题目和答案内容到右侧文本框
    def on_history_select(self, event):
        selection = self.history_listbox.curselection()
        if not selection: return
        index = selection[0]
        exam_id = self.exam_mapping.get(index)
        if not exam_id: return
        exam_data = self.db.get_exam_by_id(exam_id)
        if exam_data:
            self.question_text.delete(1.0, tk.END)
            self.answer_text.delete(1.0, tk.END)
            self.question_text.insert(1.0, exam_data[8] or "")
            self.answer_text.insert(1.0, exam_data[9] or "")
            self.current_selected_exam = exam_data

    # 打开历史试卷对应的Word文档
    def open_selected_word(self):
        if hasattr(self, 'current_selected_exam') and self.current_selected_exam:
            word_path = self.current_selected_exam[11]
            if word_path and os.path.exists(word_path):
                os.startfile(os.path.abspath(word_path))
            else:
                messagebox.showwarning("警告", "Word文档不存在或路径无效")

    # 打开历史试卷对应的HTML文件（使用默认浏览器）
    def open_selected_html(self):
        if hasattr(self, 'current_selected_exam') and self.current_selected_exam:
            html_path = self.current_selected_exam[10]
            if html_path and os.path.exists(html_path):
                webbrowser.open('file://' + os.path.abspath(html_path))
            else:
                messagebox.showwarning("警告", "HTML文件不存在或路径无效")

    # 删除选中的历史试卷记录（仅删除数据库记录，不删除文件）
    def delete_selected_exam(self):
        selection = self.history_listbox.curselection()
        if not selection:
            messagebox.showwarning("警告", "请先选择要删除的试卷")
            return
        index = selection[0]
        exam_id = self.exam_mapping.get(index)
        if not exam_id: return
        if messagebox.askyesno("确认删除", "确定要删除选中的试卷吗？"):
            self.db.delete_exam(exam_id)
            self.refresh_history()
            self.question_text.delete(1.0, tk.END)
            self.answer_text.delete(1.0, tk.END)

# 调用Coze API的流式请求函数，支持实时显示进度并拼接完整返回内容
def call_agent_stream_gui(prompt: str, bot_id: str, timeout: int, gui_instance=None, agent_name="智能体") -> str:
    headers = {"Authorization": f"Bearer {AGENT_TOKEN}", "Content-Type": "application/json"}
    data = {
        "bot_id": bot_id,
        "user_id": USER_ID,
        "stream": True,
        "auto_save_history": False,
        "additional_messages": [{"role": "user", "content": prompt, "content_type": "text"}]
    }
    try:
        if gui_instance:
            gui_instance.log_message(f"调用{agent_name}，正在接收并解析内容...")
        with requests.post(AGENT_API_URL, json=data, headers=headers, timeout=timeout, stream=True) as response:
            response.raise_for_status()
            full_content = ""
            start_time = time.time()
            last_update_time = start_time
            for line in response.iter_lines():
                if not line: continue
                try:
                    line = line.decode("utf-8")
                except UnicodeDecodeError:
                    line = line.decode("gbk", errors="ignore")
                if any(keyword in line for keyword in ["knowledge_recall", "chunks", "slice", "dataset", "1600è¯"]):
                    continue
                current_chunk = ""
                if line.startswith("data:"):
                    try:
                        data_str = line.split(":", 1)[1].strip()
                        data_json = json.loads(data_str)
                        if data_json.get("type") == "answer":
                            current_chunk = data_json.get("content", "").replace("\\\\n", "\n").replace("\\\\", "\\")
                    except:
                        if "content" in line and "\"answer\"" in line:
                            content_start = line.find("\"content\":\"") + len("\"content\":\"")
                            content_end = line.rfind("\"")
                            if content_start < content_end:
                                current_chunk = line[content_start:content_end].replace("\\\\n", "\n").replace("\\\\", "\\")
                elif line.startswith(" D") or line.startswith(".") or line.startswith(" purpose"):
                    current_chunk = line.strip() + " "
                if current_chunk:
                    full_content += current_chunk
                current_time = time.time()
                if current_time - last_update_time >= 0.5:
                    elapsed = int(current_time - start_time)
                    progress_msg = f"{agent_name}已生成{len(full_content)}个字符 {elapsed}秒"
                    if gui_instance:
                        gui_instance.update_status(progress_msg)
                    last_update_time = current_time
            elapsed_total = int(time.time() - start_time)
            if gui_instance:
                gui_instance.log_message(f"{agent_name}内容生成完成，共{len(full_content)}字符，耗时{elapsed_total}秒")
            full_content = "\n".join([line.strip() for line in full_content.split("\n") if line.strip()])
            return full_content
    except RequestException as e:
        error_msg = f"\n{agent_name}请求失败：{str(e)}"
        if gui_instance:
            gui_instance.log_message(error_msg)
        return ""

# 从返回的文本中提取指定标记的代码块内容（如```题目 或 ```答案）
def extract_code_block_content(content: str, block_start_marker: str) -> str:
    start_tag = f"```{block_start_marker}"
    end_tag = "```"
    start_idx = content.find(start_tag)
    if start_idx == -1: return ""
    content_after_start = content[start_idx + len(start_tag):]
    end_idx = content_after_start.find(end_tag)
    if end_idx == -1: return ""
    return content_after_start[:end_idx].strip()

# 检查并自动安装生成Word所需的python-docx包
def check_and_install_dependencies():
    required_packages = ['python-docx']
    missing_packages = []
    for package in required_packages:
        try:
            if package == 'python-docx':
                from docx import Document
                from docx.shared import Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            missing_packages.append(package)
    if missing_packages:
        for package in missing_packages:
            try:
                subprocess.check_call([sys.executable, "-m", "pip", "install", package])
            except:
                return False
    return True

# 将Markdown格式的题目和答案转换为带样式的HTML内容
def format_content_to_html(content: str) -> str:
    html_lines = []
    lines = content.split("\n")
    in_table = False
    table_rows = []
    for line in lines:
        line = line.strip()
        if not line:
            if in_table and table_rows:
                html_lines.append(render_table(table_rows))
                table_rows = []
                in_table = False
            html_lines.append("<br>")
            continue
        if line.startswith("# ") and not line.startswith("###"):
            html_lines.append(f"<h1 style='color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 15px; text-align: center; margin: 30px 0 20px 0;'>{line.lstrip('# ').strip()}</h1>")
        elif line.startswith("## "):
            html_lines.append(f"<h2 style='color: #2980b9; border-left: 4px solid #3498db; padding-left: 15px; margin: 25px 0 15px 0;'>{line.lstrip('# ').strip()}</h2>")
        elif line.startswith("###"):
            html_lines.append(f"<h3 style='color: #16a085; margin: 20px 0 10px 0; padding: 8px 0; border-bottom: 1px solid #ecf0f1;'>{line.lstrip('# ').strip()}</h3>")
        elif line == "---":
            html_lines.append("<hr style='border: 2px dashed #bdc3c7; margin: 30px 0;'>")
        elif line.startswith("|") and "|" in line[1:]:
            if not in_table: in_table = True
            table_rows.append(line)
        else:
            if in_table and table_rows:
                html_lines.append(render_table(table_rows))
                table_rows = []
                in_table = False
            if line.startswith("A.") or line.startswith("B.") or line.startswith("C.") or line.startswith("D."):
                html_lines.append(f"<p style='margin: 5px 0 5px 30px;'>{line}</p>")
            else:
                html_lines.append(f"<p style='margin: 10px 0; line-height: 1.6;'>{line}</p>")
    if in_table and table_rows:
        html_lines.append(render_table(table_rows))
    return "\n".join(html_lines)

# 将Markdown表格行转换为HTML表格字符串
def render_table(table_rows: list) -> str:
    if not table_rows: return ""
    html_table = ['<table style="width: 100%; border-collapse: collapse; margin: 15px 0; border: 1px solid #ddd;">']
    for i, row in enumerate(table_rows):
        row = row.strip()
        if not row.startswith('|') or not row.endswith('|'): continue
        cells = [cell.strip() for cell in row[1:-1].split('|')]
        is_header = i == 0 or any('---' in cell for cell in cells)
        tag = 'th' if is_header else 'td'
        style = 'background-color: #2c3e50; color: white; font-weight: bold; padding: 12px; border: 1px solid #ddd;' if is_header else 'padding: 10px; border: 1px solid #ddd;'
        html_cells = []
        for cell in cells:
            if '---' in cell: continue
            html_cells.append(f'<{tag} style="{style}">{cell}</{tag}>')
        if html_cells:
            html_table.append(f'<tr>{"".join(html_cells)}</tr>')
    html_table.append('</table>')
    return '\n'.join(html_table)

# 将题目和答案的Markdown内容转换为标准的Word文档（docx）
def markdown_to_word(question_content: str, answer_content: str, word_path: str) -> bool:
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        title = doc.add_heading('英语综合能力测试试卷', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if question_content:
            doc.add_heading('试卷题目', level=1)
            lines = question_content.split('\n')
            in_table = False
            table_rows = []
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue
                if line.startswith('# ') and not line.startswith('###'):
                    doc.add_heading(line.lstrip('# ').strip(), level=1)
                elif line.startswith('## '):
                    doc.add_heading(line.lstrip('# ').strip(), level=2)
                elif line.startswith('### '):
                    doc.add_heading(line.lstrip('# ').strip(), level=3)
                elif line == '---':
                    p = doc.add_paragraph()
                    p.add_run('─' * 50).bold = True
                elif line.startswith('|') and '|' in line[1:]:
                    if not in_table:
                        in_table = True
                        table_rows = []
                    table_rows.append(line)
                else:
                    if in_table and table_rows:
                        _add_table_to_doc(doc, table_rows)
                        table_rows = []
                        in_table = False
                    if line.startswith(('A.', 'B.', 'C.', 'D.')):
                        p = doc.add_paragraph()
                        p.add_run(line).font.size = Inches(0.14)
                    else:
                        p = doc.add_paragraph(line)
                        p.style.font.size = Inches(0.14)
            if in_table and table_rows:
                _add_table_to_doc(doc, table_rows)
        doc.add_page_break()
        if answer_content:
            doc.add_heading('参考答案与解析', level=1)
            lines = answer_content.split('\n')
            in_table = False
            table_rows = []
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue
                if line.startswith('# ') and not line.startswith('###'):
                    doc.add_heading(line.lstrip('# ').strip(), level=1)
                elif line.startswith('## '):
                    doc.add_heading(line.lstrip('# ').strip(), level=2)
                elif line.startswith('### '):
                    doc.add_heading(line.lstrip('# ').strip(), level=3)
                elif line == '---':
                    p = doc.add_paragraph()
                    p.add_run('─' * 50).bold = True
                elif line.startswith('|') and '|' in line[1:]:
                    if not in_table:
                        in_table = True
                        table_rows = []
                    table_rows.append(line)
                else:
                    if in_table and table_rows:
                        _add_table_to_doc(doc, table_rows)
                        table_rows = []
                        in_table = False
                    if line.startswith(('A.', 'B.', 'C.', 'D.')):
                        p = doc.add_paragraph()
                        p.add_run(line).font.size = Inches(0.14)
                    else:
                        p = doc.add_paragraph(line)
                        p.style.font.size = Inches(0.14)
            if in_table and table_rows:
                _add_table_to_doc(doc, table_rows)
        doc.save(word_path)
        return True
    except Exception as e:
        print(f"生成Word文档失败: {str(e)}")
        return False

# 辅助函数：将Markdown表格行添加到docx文档的表格中
def _add_table_to_doc(doc, table_rows):
    from docx.enum.table import WD_TABLE_ALIGNMENT
    if not table_rows: return
    num_cols = len(table_rows[0].split('|')) - 1
    num_rows = len([row for row in table_rows if not any('---' in cell for cell in row.split('|'))])
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    row_idx = 0
    for i, row in enumerate(table_rows):
        if any('---' in cell for cell in row.split('|')): continue
        cells = [cell.strip() for cell in row[1:-1].split('|')]
        for col_idx, cell_content in enumerate(cells):
            if col_idx < num_cols:
                table.cell(row_idx, col_idx).text = cell_content
        row_idx += 1

# 程序入口：创建主窗口并启动GUI事件循环
if __name__ == "__main__":
    root = tk.Tk()
    app = EnglishExamGeneratorGUI(root)
    root.mainloop()